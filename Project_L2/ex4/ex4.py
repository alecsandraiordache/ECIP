# Make a series of experiments in which the population number starts from 
# a) 0.1 b) 0.5 c) 1.0 d) 2
# Make a report in which you note the observations for r = 4

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os


R = 4
initial_populations = [0.1, 0.5, 1.0, 2.0]
generations = 20


fig, axes = plt.subplots(2, 2, figsize=(10, 8))
fig.suptitle(f"Population Growth Experiments (R = {R})", fontsize=16)

axes = axes.flatten()

for i, P0 in enumerate(initial_populations):
    population_history = [P0]
    current_P = P0
    
    for _ in range(generations):
        try:
            current_P = R * current_P * (1 - current_P)
            population_history.append(current_P)
        except OverflowError:
            population_history.append(float('-inf'))
            break 
            
        
    ax = axes[i]
    ax.plot(range(len(population_history)), population_history, marker='o', color='teal')
    titles = ["a) P0 = 0.1", "b) P0 = 0.5", "c) P0 = 1.0", "d) P0 = 2.0"]
    ax.set_title(titles[i])
    ax.set_xlabel("Generations")
    ax.set_ylabel("Population Fraction")
    ax.grid(True, linestyle='--', alpha=0.6)
    ax.axhline(0, color='red', linestyle='--', linewidth=1) 

plt.tight_layout(rect=[0, 0.03, 1, 0.95])


image_filename = 'experiment_results_R4.png'
plt.savefig(image_filename, dpi=150)
plt.close() 
print("Chart generated and saved temporarily.")


doc = Document()


doc.add_heading('Observation Report: Fate of the Population for R = 4', 0)


doc.add_paragraph(
    "This report analyzes the discrete logistic growth model under extreme conditions. "
    "When the growth factor is set to R = 4, the environment is at its maximum chaotic limit. "
    "The fate of the population depends entirely on its starting size (P0)."
)


doc.add_heading('1. Experimental Charts', level=1)
doc.add_picture(image_filename, width=Inches(6.0))


doc.add_heading('2. Detailed Observations', level=1)


doc.add_heading('a) Starting Population: P0 = 0.1 (10% capacity)', level=2)
doc.add_paragraph(
    "FULL CHAOS. The population quickly shoots up, then fluctuates wildly between 0 and 1. "
    "It never stabilizes and never repeats the same pattern. It constantly bounces between "
    "near-extinction and overpopulation."
)


doc.add_heading('b) Starting Population: P0 = 0.5 (50% capacity)', level=2)
doc.add_paragraph(
    "DELAYED EXTINCTION. Because R=4 is so high, starting at exactly 50% causes the population "
    "to hit exactly 1.0 (100% capacity) in generation 1. Having consumed absolutely all resources, "
    "it crashes to exactly 0.0 in generation 2. Extinction is permanent."
)


doc.add_heading('c) Starting Population: P0 = 1.0 (100% capacity)', level=2)
doc.add_paragraph(
    "INSTANT EXTINCTION. The population begins already at the absolute maximum limit of the "
    "environment. There are zero resources left for reproduction. It instantly drops to 0.0 "
    "in generation 1."
)


doc.add_heading('d) Starting Population: P0 = 2.0 (200% capacity)', level=2)
doc.add_paragraph(
    "MATHEMATICAL COLLAPSE. The population starts at double what the environment can support. "
    "In the formula, the (1 - 2.0) term creates a negative multiplier (-1). The population "
    "instantly drops into impossible negative numbers and spirals rapidly into negative infinity. "
    "Biologically, the environment was so severely over-consumed that total ecosystem collapse "
    "occurred instantly."
)

doc_filename = 'Observation_Report_R4.docx'
doc.save(doc_filename)

if os.path.exists(image_filename):
    os.remove(image_filename)
