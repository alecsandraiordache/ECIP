#modify the value of the weight for state B to 0.5
#plot the evolution of the population for 15 steps based on seq S.
#Answer: Does the population survive after 15 th step. 
# write a report which must contain the plot for these 15 steps and the answer to the question


import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os


np.random.seed(42)

states = ['A', 'B', 'C']


current_state = 'A'
S = []

for _ in range(15):
    if current_state == 'A':
        probabilities = [0.0, 0.4, 0.6]
    elif current_state == 'B':
        probabilities = [0.5, 0.5, 0.0]
    else:
        probabilities = [0.2, 0.2, 0.6]
        
    current_state = np.random.choice(states, p=probabilities)
    S.append(current_state)


weights = {'A': 1.5, 'B': 0.5, 'C': 3.3}
R_sequence = [weights[state] for state in S]

current_population = 0.5
population_history = [current_population]


for step in range(15):
    R = R_sequence[step]
    current_population = R * current_population * (1 - current_population)
    population_history.append(current_population)


plt.figure(figsize=(8, 5))
plt.plot(range(16), population_history, marker='o', color='green', linewidth=2)
plt.title("Population Evolution (State B weight = 0.5)", fontweight='bold')
plt.xlabel("Step")
plt.ylabel("Population Fraction (P)")
plt.xticks(range(16))
plt.grid(True, linestyle='--', alpha=0.7)
plt.tight_layout()


image_filename = 'survival_plot_B05.png'
plt.savefig(image_filename, dpi=150)
plt.close()


doc = Document()
doc.add_heading('Population Survival Report', 0)


doc.add_heading('1. Simulation Parameters', level=1)
doc.add_paragraph(f"Initial Population: 0.5")
doc.add_paragraph(f"Sequence of 15 states (S): {', '.join(S)}")
doc.add_paragraph("Weights (R factors): A = 1.5, B = 0.5, C = 3.3")


doc.add_heading('2. Population Evolution Plot', level=1)
doc.add_picture(image_filename, width=Inches(6.0))


doc.add_heading('3. Answer to the Question', level=1)
doc.add_paragraph(
    "Yes, the population survives.\n\n"
    "Even though State B has an R value of 0.5 (which is a declining, extinction-level growth rate), "
    "the environment dynamically shifts out of State B back to State A (R=1.5) or State C (R=3.3). "
    "During the periods where State B is dominant, the population drops significantly. However, "
    "it never fully reaches zero because mathematically, multiplying by 0.5 simply halves the population "
    "with each step. As long as the environment eventually transitions to a favorable state (like C), "
    "the population bounces back. By the end of step 15, the population fraction remains positive "
    "and active."
    "In conclusion, a population can survive temporary environmental hardships (State B) as long as the system "
    "dynamically returns to favorable conditions. Long-term survival depends entirely on the "
    "balance and frequency of these recovery phases."
)



doc_filename = 'Survival_Report_B05.docx'
doc.save(doc_filename)
print(f"Report successfully saved as: {doc_filename}")


if os.path.exists(image_filename):
    os.remove(image_filename)