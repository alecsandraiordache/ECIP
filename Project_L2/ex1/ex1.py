#make a simulation by using the logistic regresion expressiion( = population  growth)
#use a growth factor R equal to 0.9, make a raport (word doc) wich describes the fate of
#population for R = 0.9

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def simulate_logistic_growth(R, P0, generations):
    """Simulates population using the logistic map."""
    populations = [P0]
    for _ in range(generations):
        P_next = R * populations[-1] * (1 - populations[-1])
        populations.append(P_next)
    return populations

R = 0.9           
P0 = 0.5         
generations = 50

population_data = simulate_logistic_growth(R, P0, generations)


plt.figure(figsize=(8, 5))
plt.plot(range(generations + 1), population_data, marker='o', color='red', linestyle='-')
plt.title(f'Logistic Population Growth (R = {R})')
plt.xlabel('Generation (n)')
plt.ylabel('Population Fraction (P)')
plt.grid(True)
plt.tight_layout()


image_filename = 'population_simulation.png'
plt.savefig(image_filename)
print("Simulation complete. Chart saved.")


doc = Document()


doc.add_heading('Population Growth Simulation Report', 0)


doc.add_heading('1. The Model', level=1)
doc.add_paragraph(
    "This report simulates population growth using the discrete logistic growth model "
    "(logistic map), defined by the formula: P(n+1) = R * P(n) * (1 - P(n)). "
    "Here, P represents the population as a fraction of the environment's maximum "
    "carrying capacity, and R represents the growth factor."
)


doc.add_heading('2. Simulation Parameters', level=1)
doc.add_paragraph(f"• Growth Factor (R): {R}")
doc.add_paragraph(f"• Initial Population (P0): {P0}")
doc.add_paragraph(f"• Generations Simulated: {generations}")


doc.add_heading('3. The Fate of the Population', level=1)
doc.add_paragraph(
    "When the growth factor R is strictly less than 1 (R = 0.9), the population is "
    "doomed to extinction. Because the reproduction rate is too low to replace the "
    "current generation—even without factoring in the dampening effect of the carrying "
    "capacity (1 - P)—the population steadily declines. As shown in the graph below, "
    "the population fraction rapidly approaches zero over successive generations."
)

doc.add_heading('4. Simulation Chart', level=1)
doc.add_picture(image_filename, width=Inches(6))


doc_filename = 'Population_Report_R0.9.docx'
doc.save(doc_filename)
print(f"Report successfully saved as: {doc_filename}")