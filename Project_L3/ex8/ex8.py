#compute average of noisy measurements
#compare estimation with true value
#discuss estimation eroor on word file 

import numpy as np
from docx import Document

n_samples = 10000
true_x = 10.0
np.random.seed(42)

noise = np.random.normal(loc=0.0, scale=2.0, size=n_samples)
measurements_y = true_x + noise

estimated_x = np.mean(measurements_y)
estimation_error = abs(true_x - estimated_x)

doc = Document()

doc.add_heading('Discussion: Value Estimation and Estimation Error', 0)

doc.add_heading('1. Computing the Average of Noisy Measurements', level=2)
doc.add_paragraph(
    f"In a real-world scenario, the true value (x = {true_x}) of a system is unknown. "
    "All we have at our disposal is a set of imperfect measurements affected by noise (y = x + noise). "
    "To find the real value, we use statistical estimation. The simplest and most effective method is computing "
    f"the arithmetic mean of all the obtained measurements. In our simulation, the average of the {n_samples} "
    f"noisy measurements yielded an estimated value of approximately {estimated_x:.4f}."
)

doc.add_heading('2. Comparing the Estimation with the True Value', level=2)
doc.add_paragraph(f"• True Value (x): {true_x:.4f}\n• Estimated Value (x_hat): {estimated_x:.4f}")
doc.add_paragraph(
    "It can be observed that the estimated value is extremely close to the real value, even though each individual "
    "measurement was corrupted by noise with a relatively large standard deviation (±2.0 deviation)."
)


doc.add_heading('3. Estimation Error', level=2)
doc.add_paragraph(
    "The estimation error represents the absolute difference between the estimated value and the true value "
    f"(Error = |x - x_hat|). In this specific case, the error is only {estimation_error:.4f}."
)
doc.add_paragraph(
    "Why does this work so well? This high degree of accuracy is due to the properties of Gaussian noise. "
    "Because the noise has a mean of zero (zero-mean), the positive and negative errors cancel each other out "
    "when we average a large number of samples, entirely consistent with the Law of Large Numbers."
)


file_name = 'Estimation_Error_Discussion.docx'
doc.save(file_name)

print("-" * 50)
print(f"DONE! The file '{file_name}' has been successfully saved in your current directory.")
print("-" * 50)